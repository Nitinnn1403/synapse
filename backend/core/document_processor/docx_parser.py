import os
from docx import Document as DocxDocument

from core.document_processor.base import (
    BaseParser, ParsedDocument, PageContent, ExtractedTable
)


class DocxParser(BaseParser):
    async def parse(self, file_path: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        full_text_parts: list[str] = []
        tables: list[ExtractedTable] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                # Convert to markdown table
                header = "| " + " | ".join(rows[0]) + " |"
                separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
                body_rows = ["| " + " | ".join(row) + " |" for row in rows[1:]]
                markdown = "\n".join([header, separator] + body_rows)
                tables.append(ExtractedTable(markdown=markdown))
                full_text_parts.append(markdown)

        full_text = "\n\n".join(full_text_parts)

        # Treat entire docx as a single "page"
        page = PageContent(
            page_number=1,
            text=full_text,
            tables=tables,
        )

        return ParsedDocument(
            text=full_text,
            pages=[page],
            metadata={"filename": os.path.basename(file_path)},
            tables=tables,
        )
